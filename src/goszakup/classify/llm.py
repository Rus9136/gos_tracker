"""LLM-классификатор: читает PDF тех.задания и заполняет LotAnalysis.

Это ТРЕТЬЯ ступень классификации поверх enstru/keyword pre-filter из it.py.
Вызывается ОДИН раз при парсинге (или по команде reanalyze) — никаких
LLM-вызовов на запросах UI.

Идемпотентность: пропускаем переанализ, если у лота уже есть запись
LotAnalysis с тем же analyzer_version И тем же tz_sha256.
"""

from __future__ import annotations

import json
import logging
import os
import re
from collections.abc import Iterable
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta, timezone
from functools import lru_cache
from typing import Literal

from pydantic import BaseModel, Field, ValidationError
from sqlalchemy import select
from sqlalchemy.orm import Session, object_session

from ..db.models import Announcement, Document, Lot, LotAnalysis, PlanPoint
from ..jobs.plans import plan_root_from_number
from ..scraper.katos import region_name
from ..scraper.modal_files import is_tz_like_name
from ..watchlist import should_analyze
from .extractive_summary import extract_summary
from .it import classify as _it_subcategory
from .rules import CONFIDENCE_THRESHOLD, RULES_VERSION, classify_lot
from .simhash import (
    HAMMING_THRESHOLD,
    from_signed64,
    hamming,
    to_signed64,
)
from .simhash import (
    simhash as compute_simhash,
)
from .usage import record_call, usage_from_response

log = logging.getLogger(__name__)

# Версия меняется при правках промпта, схемы AnalysisResult или провайдера —
# старые записи будут переанализированы при следующем запуске.
ANALYZER_VERSION = "llm-v4-gpt-oss-120b-brief"

# Дефолт — gpt-oss-120b на Cerebras: open-weights модель с быстрым инференсом,
# поддерживает tool calling со `strict: True` (constrained decoding —
# гарантия валидной схемы). Перебивается ENV GZ_LLM_MODEL.
DEFAULT_MODEL = "gpt-oss-120b"

# Сколько символов текста ТЗ отдаём модели. Кириллица у gpt-oss ≈ 3 символа
# на токен, контекст модели на Cerebras — 64K токенов: 60K символов ≈ 20K
# токенов, с запасом остаётся место под промпт, историю чата и ответ. На
# прежнем пороге 30K обрезалось каждое пятое ТЗ (замер 08.2026), причём
# молча — модель не знала, что видит не весь документ.
MAX_TZ_CHARS = 60_000

# Минимальный непробельный размер «осмысленного» PDF-текста. Меньше — считаем,
# что текстового слоя нет (битый PDF / скан), и идём по low-confidence пути.
MIN_TEXT_CHARS = 200


# === Pydantic-схема ответа модели ===

DevCategory = Literal[
    "1c_development",
    "web_development",
    "mobile_dev",
    "integration",
    "software_support",
    "it_infra",
    "hardware",
    "not_dev",
]
VendorLockRisk = Literal["low", "medium", "high"]
Confidence = Literal["high", "low"]


class AnalysisResult(BaseModel):
    dev_category: DevCategory = Field(description="Главная категория")
    tech_stack: list[str] = Field(
        default_factory=list,
        description="Технологии/платформы (1С, FastAPI, React и т.п.). Пусто, если не ясно",
    )
    tz_summary: str = Field(
        description=(
            "СТРОГО НА РУССКОМ ЯЗЫКЕ. Бриф для участника закупки, 4-7 "
            "предложений: что закупают и в каком объёме, ключевые требования "
            "(технические и к участнику), сроки и условия оплаты, на что "
            "обратить внимание. Если исходный ТЗ на казахском/английском — "
            "переведи на русский."
        )
    )
    solo_feasible: bool = Field(description="Реально ли выполнить одному разработчику")
    vendor_lock_risk: VendorLockRisk = Field(description="Риск заточки под конкретного поставщика")
    analysis_confidence: Confidence = Field(
        description="high — классифицировано по тексту ТЗ; low — без ТЗ"
    )


# === Промпт ===

SYSTEM_PROMPT = """Ты — аналитик тендеров на стороне участника гос.закупок Казахстана.
Тебе дан лот и текст технической спецификации (ТЗ). Верни СТРОГО
структурированный результат через вызов инструмента `submit_classification`.
Без преамбулы, без markdown, без текста вне инструмента.

ВАЖНО: ВСЕ строковые поля — ТОЛЬКО на русском языке. Если ТЗ написано на
казахском или английском — переводи на русский. Технические термины (1С,
FastAPI, PostgreSQL, ERP) и имена собственные можно оставлять как есть.

=== tz_summary — главное поле. Это бриф, по которому участник за 20 секунд
решает, стоит ли открывать документы. Пиши так, чтобы из брифа было понятно
то, чего НЕ видно из названия лота. 4-7 предложений, одним абзацем, без
списков и заголовков. Порядок:

1. Что закупают и сколько: предмет, объём/количество, место поставки или
   выполнения работ. Не повторяй название лота дословно — раскрой его.
2. Ключевые требования: 2-4 самых значимых технических требования (модель,
   характеристики, функциональность, объём работ) и требования к
   участнику (опыт, лицензии, сертификаты, персонал), если они есть в ТЗ.
3. Сроки и условия: срок поставки/выполнения, гарантия, этапность, аванс
   или порядок оплаты — только если это есть в тексте.
4. На что обратить внимание: нетипичные или жёсткие условия, штрафы,
   признаки заточки под конкретного поставщика, противоречия в ТЗ.

Не пиши общих фраз вида «в соответствии с техническим заданием заказчика»
и «согласно требованиям» — они не несут информации. Не выдумывай: если
сроков или требований к участнику в ТЗ нет — просто не упоминай их. Если
ТЗ недоступно — бриф строится по названию, ENSTRU и полям объявления, и
это должно быть видно из текста («ТЗ не приложено, по названию …»).

=== dev_category — выбери РОВНО ОДНУ. Категории описывают IT-разработку;
для лота вне ИТ это `not_dev`, и это нормальный, ожидаемый ответ.

- `1c_development` — разработка / доработка / внедрение / конфигурирование
  на платформе 1С (1С:Предприятие, БП, ЗУП и т.п.). Услуги программиста 1С.
  НЕ путать: продление лицензии 1С, поставка коробки 1С, абонентское
  сопровождение готовой 1С без разработки → `software_support`.
- `web_development` — сайт, портал, веб-приложение, веб-интерфейс, CMS,
  e-commerce. Создание/редизайн/доработка веб-системы.
- `mobile_dev` — нативное или кроссплатформенное мобильное приложение
  (iOS/Android/Flutter/RN).
- `integration` — интеграция систем, обмен данными между системами, ETL,
  API-шлюзы между существующими ИС, шины данных.
- `software_support` — техническая поддержка / сопровождение готового ПО,
  продление лицензий, поставка коробочного ПО без разработки, абонентское
  обслуживание ИС без новой функциональности, антивирус.
- `it_infra` — администрирование сетей, сетевая безопасность,
  программно-аппаратные комплексы (ПАК), видеонаблюдение, СКУД, ЦОД,
  серверная инфраструктура как услуга.
- `hardware` — поставка ИТ-железа: ПК, ноутбуки, серверы, принтеры,
  коммутаторы, мониторы, расходники к ним, коробочные лицензии.
  Приборы, оборудование и техника вне ИТ (медицинская, измерительная,
  промышленная, мебель, транспорт) — это `not_dev`, а не `hardware`.
- `not_dev` — всё, что не разработка/поддержка/ИТ-железо:
  * «проектно-сметная документация» на вентиляцию, канализацию, ремонт,
    электросеть, благоустройство — это СТРОИТЕЛЬНОЕ проектирование, не софт;
  * консультации/обучение/семинары вне ИТ;
  * товары, работы и услуги любых других отраслей.

=== vendor_lock_risk — риск заточки под единственного поставщика. `high`:
- размытое название («Согласно технической документации», «Обеспечение
  финансового процесса»), без конкретики;
- требования, которые объективно может выполнить только один поставщик
  (упоминание конкретной существующей системы заказчика, патентованного
  решения, эксклюзивной лицензии, конкретной торговой марки без «или
  эквивалент»);
- нетипично короткие сроки (несколько дней на сложную работу или поставку);
- узкоспецифичные технические отсылки к существующей инфраструктуре.

=== analysis_confidence: есть текст ТЗ → `high`; ТЗ недоступно
(пометка «ТЗ: недоступно») → классифицируй по названию + ENSTRU + способу
закупки и ОБЯЗАТЕЛЬНО верни `low`.

=== solo_feasible: `true` — задача реалистична для одного исполнителя
(до ~3 месяцев работы соло, без лицензируемой деятельности и без крупной
поставки). `false` — нужна команда, лицензированная организация, либо
физическое присутствие/поставка в масштабе.

=== tech_stack — массив коротких строк («1С», «React», «PostgreSQL»). Пусто,
если из текста не ясно или категория `hardware`/`not_dev`.
"""

# OpenAI-совместимый формат tools (используется Cerebras). `strict: True`
# включает constrained decoding на стороне сервера — модель физически не
# сможет вернуть невалидный JSON. `additionalProperties: false` обязательно
# для strict-режима.
CLASSIFICATION_TOOL = {
    "type": "function",
    "function": {
        "name": "submit_classification",
        "strict": True,
        "description": "Сохранить классификацию лота. Вызови РОВНО ОДИН РАЗ.",
        "parameters": {
            "type": "object",
            "additionalProperties": False,
            "properties": {
                "dev_category": {
                    "type": "string",
                    "enum": [
                        "1c_development",
                        "web_development",
                        "mobile_dev",
                        "integration",
                        "software_support",
                        "it_infra",
                        "hardware",
                        "not_dev",
                    ],
                },
                "tech_stack": {
                    "type": "array",
                    "items": {"type": "string"},
                },
                "tz_summary": {
                    "type": "string",
                    "description": (
                        "Бриф для участника, 4-7 предложений одним абзацем: что "
                        "закупают и в каком объёме, ключевые технические "
                        "требования и требования к участнику, сроки и условия "
                        "оплаты, на что обратить внимание. Без общих фраз, без "
                        "выдумок. ТОЛЬКО на русском языке — даже если исходный "
                        "ТЗ на казахском/английском, переводи на русский."
                    ),
                },
                "solo_feasible": {"type": "boolean"},
                "vendor_lock_risk": {
                    "type": "string",
                    "enum": ["low", "medium", "high"],
                },
                "analysis_confidence": {
                    "type": "string",
                    "enum": ["high", "low"],
                },
            },
            "required": [
                "dev_category",
                "tech_stack",
                "tz_summary",
                "solo_feasible",
                "vendor_lock_risk",
                "analysis_confidence",
            ],
        },
    },
}


# === Выбор «ТЗ»-документа из приложений к объявлению ===

# Предикат `is_tz_like_name` живёт в scraper/modal_files.py — там же, где
# принимается решение «разворачивать ли modal у этой строки». Здесь его
# переиспользуем для финального выбора из уже скачанных файлов.

# PDF предпочтительнее DOCX: на goszakup PDF — это подписанный оригинал,
# DOCX обычно шаблоны/приложения. Если есть оба похожих — берём PDF.
_PREFERRED_EXTS = (".pdf", ".docx")


def _ext(path: str | None) -> str:
    return (path or "").lower().rsplit(".", 1)[-1] if path and "." in path else ""


def _detect_format(path: str | None) -> str:
    """'pdf' / 'docx' / 'doc' / ''. Расширение первично; без него — магические
    байты: API-путь до фикса suggested_name сохранял файлы под хешем из URL
    без расширения, и такие лежат на диске тысячами."""
    ext = _ext(path)
    if ext.isalnum() and len(ext) <= 5:
        return ext
    if not path:
        return ""
    try:
        with open(path, "rb") as fh:
            head = fh.read(4)
    except OSError:
        return ""
    if head == b"%PDF":
        return "pdf"
    if head == b"PK\x03\x04":
        # zip-контейнер — docx только если внутри word/document.xml
        # (xlsx/обычный zip нам не нужны).
        import zipfile

        try:
            with zipfile.ZipFile(path) as zf:
                if any(n.startswith("word/") for n in zf.namelist()):
                    return "docx"
        except Exception:
            return ""
    return ""


def pick_tz_document(documents: Iterable[Document]) -> Document | None:
    """Финальный выбор файла для подачи в LLM: ТЗ, со скачанным телом."""
    formats: dict[int, str] = {}
    downloaded = []
    for d in documents:
        if not (d.local_path and d.sha256):
            continue
        fmt = _detect_format(d.local_path)
        if fmt in {"pdf", "docx"}:
            formats[id(d)] = fmt
            downloaded.append(d)
    if not downloaded:
        return None

    def name_haystack(d: Document) -> str:
        # Имя файла на диске тоже учитываем — там бывают маркеры вида techspec_…
        from pathlib import Path

        fname = Path(d.local_path).name if d.local_path else ""
        return f"{d.name or ''} {d.attribute or ''} {fname}"

    matches = [d for d in downloaded if is_tz_like_name(name_haystack(d))]
    candidates = matches or downloaded

    # Узкая «техническая спецификация» точнее, чем «конкурсная документация»
    # (последняя — шаблон с приложениями + ТЗ внутри; нам нужен ТЗ).
    _TECHSPEC_RE = re.compile(
        r"техническ\w*\s*специф|\btechspec_|тех(ническое)?\s*задани",
        re.IGNORECASE,
    )

    def specificity(d: Document) -> int:
        from pathlib import Path
        fname = Path(d.local_path).name if d.local_path else ""
        return 0 if _TECHSPEC_RE.search(f"{d.name or ''} {fname}") else 1

    # PDF > DOCX; крупнее > мельче (внутри той же специфичности).
    def sort_key(d: Document) -> tuple[int, int, int]:
        ext_score = 0 if formats[id(d)] == "pdf" else 1
        return (specificity(d), ext_score, -(d.size or 0))

    candidates.sort(key=sort_key)
    return candidates[0]


# === Извлечение текста: PDF и DOCX через единый диспатчер ===


def extract_text(path: str) -> str | None:
    """Возвращает текст документа или None, если содержимого нет / формат не поддержан."""
    ext = _detect_format(path)
    if ext == "pdf":
        text = _extract_pdf_text(path)
    elif ext == "docx":
        text = _extract_docx_text(path)
    elif ext == "doc":
        # Legacy .doc — для извлечения нужен LibreOffice/antiword; не тащим
        # зависимость, просто пропускаем (lot уйдёт в low-confidence путь).
        log.warning("формат .doc не поддерживается (%s), пропускаю", path)
        return None
    else:
        return None
    if text is None:
        return None
    meaningful = sum(1 for ch in text if not ch.isspace())
    if meaningful < MIN_TEXT_CHARS:
        return None
    if len(text) > MAX_TZ_CHARS:
        total = len(text)
        text = (
            text[:MAX_TZ_CHARS]
            + f"\n[… текст ТЗ обрезан: показаны первые {MAX_TZ_CHARS} из {total} символов]"
        )
    return text


def _extract_pdf_text(path: str) -> str | None:
    # OCR не запускается: продукт ориентирован на PDF с нормальным текстовым
    # слоем. Сканы без слоя → low-confidence путь без ТЗ.
    try:
        import pdfplumber
    except ImportError:
        log.warning("pdfplumber not installed; cannot extract PDF text")
        return None
    try:
        with pdfplumber.open(path) as pdf:
            chunks: list[str] = []
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t:
                    chunks.append(t)
                if sum(len(c) for c in chunks) > MAX_TZ_CHARS * 2:
                    break
        return "\n".join(chunks).strip()
    except Exception as e:
        log.warning("pdfplumber failed on %s: %s", path, e)
        return None


def _extract_docx_text(path: str) -> str | None:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        log.warning("python-docx not installed; cannot extract DOCX text")
        return None
    try:
        doc = DocxDocument(path)
    except Exception as e:
        log.warning("python-docx failed to open %s: %s", path, e)
        return None
    parts: list[str] = [p.text for p in doc.paragraphs if p.text]
    # Тех.задание часто оформлено таблицами — забираем и их.
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


# === Сборка пейлоада и вызов модели ===


def _build_user_message(
    lot: Lot, announcement: Announcement | None, tz_text: str | None
) -> str:
    lines = [
        f"Название лота: {lot.name or '—'}",
        f"ENSTRU: {lot.enstru or '—'}",
        f"Заказчик: {lot.customer.name if lot.customer else '—'}",
        f"Способ закупки: {lot.method or (announcement.method if announcement else '—')}",
        f"Плановая сумма ₸: {lot.plan_amount or '—'}",
        # Подкатегория считается на лету (classify/it.py): промпт байт-в-байт
        # как до пивота — бамп ANALYZER_VERSION не нужен.
        f"IT-категория (pre-filter): {_it_subcategory(lot.enstru, lot.name) or '—'}",
    ]
    if announcement and announcement.attributes:
        lines.append(f"Признаки объявления: {announcement.attributes}")
    if tz_text:
        lines.append("")
        lines.append("--- ТЕКСТ ТЕХНИЧЕСКОГО ЗАДАНИЯ ---")
        lines.append(tz_text)
    else:
        lines.append("")
        lines.append("ТЗ: недоступно (нет файла или текстового слоя) — классифицируй по полям выше.")
    return "\n".join(lines)


@dataclass
class _Outcome:
    """Внутренний результат: либо валидный AnalysisResult, либо ошибка."""

    result: AnalysisResult | None
    error: str | None = None
    usage: dict | None = None


# Маркеры транзиентных отказов (по типу исключения и тексту) — ретраим их.
# Детерминированные (401/403/400/invalid) сюда НЕ попадают: повтор не помогает.
_RETRYABLE_LLM_MARKERS = (
    "429", "queue_exceeded", "too_many_requests", "rate_limit", "rate limit",
    "500", "502", "503", "504",
    "internalserver", "internal server", "service unavailable", "serviceunavailable",
    "bad gateway", "gateway timeout", "overloaded", "temporarily unavailable",
    "timeout", "timed out", "connection", "connect", "econnreset", "reset by peer",
)


def _is_retryable_llm_error(e: Exception) -> bool:
    text = f"{type(e).__name__} {e}".lower()
    return any(m in text for m in _RETRYABLE_LLM_MARKERS)


def _call_llm(
    lot: Lot, announcement: Announcement | None, tz_text: str | None
) -> _Outcome:
    api_key = os.environ.get("CEREBRAS_API_KEY")
    if not api_key:
        return _Outcome(None, "CEREBRAS_API_KEY не задан")

    try:
        from cerebras.cloud.sdk import Cerebras
    except ImportError:
        return _Outcome(None, "cerebras-cloud-sdk не установлен")

    model = os.environ.get("GZ_LLM_MODEL", DEFAULT_MODEL)
    client = Cerebras(api_key=api_key)

    user_msg = _build_user_message(lot, announcement, tz_text)

    # Транзиентные отказы провайдера (429 бурст free-tier, 5xx, таймауты,
    # обрывы соединения) ретраим с бэкоффом — иначе массовый прогон теряет
    # пачку лотов на ровном месте, а восстановление возможно только вручную.
    _RETRY_DELAYS = (5.0, 15.0, 30.0)

    resp = None
    last_err: Exception | None = None
    for attempt, delay in enumerate([0.0, *_RETRY_DELAYS]):
        if delay:
            log.info("LLM retry #%d for lot %s after %.0fs", attempt, lot.id, delay)
            import time as _time
            _time.sleep(delay)
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": user_msg},
                ],
                tools=[CLASSIFICATION_TOOL],
                tool_choice={
                    "type": "function",
                    "function": {"name": "submit_classification"},
                },
                parallel_tool_calls=False,
                # Задача структурная (классификация в фиксированную схему),
                # длинные цепочки рассуждений не нужны и удорожают вызов.
                reasoning_effort="low",
                max_tokens=2048,
                timeout=60.0,
            )
            break  # успех — выходим из ретрая
        except Exception as e:
            last_err = e
            # Транзиентные — ретраим; детерминированные (401/400/невалидный
            # запрос) сразу наружу, повтор не поможет.
            if _is_retryable_llm_error(e):
                continue
            return _Outcome(None, f"Cerebras API call failed: {e}")

    if resp is None:
        return _Outcome(
            None, f"Cerebras API: транзиентный отказ после {len(_RETRY_DELAYS)} ретраев: {last_err}"
        )

    # Токены списаны вне зависимости от того, распарсился ли результат — usage
    # тащим во все ветки ниже, чтобы учёт расходов не терял неудачные вызовы.
    usage = usage_from_response(resp, model)

    choice = resp.choices[0].message if resp.choices else None
    if choice is None or not choice.tool_calls:
        finish = resp.choices[0].finish_reason if resp.choices else "—"
        return _Outcome(None, f"модель не вызвала tool: finish_reason={finish}", usage)

    tool_call = choice.tool_calls[0]
    # У Cerebras (OpenAI-формат) arguments — JSON-строка, не dict
    # (в отличие от Anthropic, который отдавал распарсенный dict).
    try:
        tool_input = json.loads(tool_call.function.arguments)
    except json.JSONDecodeError as e:
        return _Outcome(None, f"невалидный JSON в tool_call.arguments: {e}", usage)

    try:
        parsed = AnalysisResult.model_validate(tool_input)
    except ValidationError as e:
        return _Outcome(None, f"Pydantic validation failed: {e}", usage)

    # Если ТЗ нет — принудительно low, даже если модель забыла.
    if tz_text is None and parsed.analysis_confidence != "low":
        parsed = parsed.model_copy(update={"analysis_confidence": "low"})

    return _Outcome(parsed, usage=usage)


# === Точка входа из пайплайна ===


def analyze_and_save(session: Session, lot: Lot, *, force: bool = False) -> bool:
    """Возвращает True, если сделан новый LLM-вызов; False — пропуск или ошибка.

    Не бросает исключения наружу: пайплайн `daily` должен остаться рабочим
    при любой проблеме с LLM (нет ключа, нет сети, битый PDF и т.д.).

    `force=True` — обойти идемпотентность и переанализировать даже при
    совпадении `(analyzer_version, tz_sha256)`. Используется ручной кнопкой
    в UI.
    """
    try:
        return _analyze_inner(session, lot, force=force)
    except Exception as e:  # последняя линия защиты
        log.exception("analyze_and_save crashed for lot %s: %s", lot.id, e)
        return False


def _analyze_inner(session: Session, lot: Lot, *, force: bool = False) -> bool:
    # Вне watchlist LLM не вызываем — это pre-filter условие. `force` его
    # обходит: это ручная кнопка в UI, осознанное намерение человека и ровно
    # один вызов (иначе админ с пустым scope не смог бы проанализировать
    # вообще ничего — watchlist собирается по вертикалям подписчиков).
    if not force and not should_analyze(session, lot):
        return False

    announcement = lot.announcement
    documents: list[Document] = list(announcement.documents) if announcement else []

    tz_doc = pick_tz_document(documents)
    tz_sha = tz_doc.sha256 if tz_doc else None

    prev = lot.analysis
    if (
        not force
        and prev is not None
        and prev.analyzer_version in (ANALYZER_VERSION, RULES_VERSION)
        and prev.tz_sha256 == tz_sha
    ):
        return False  # уже актуально (LLM или правила на этой версии)

    # Если уже есть LLM-запись (любой LLM-версии), правила её не вытесняют —
    # идём сразу в _call_llm. Это защищает от «отката качества» при апгрейде
    # пайплайна, когда правила могли бы случайно записать поверх LLM-анализа.
    prev_is_llm = prev is not None and (prev.analyzer_version or "").startswith("llm-")

    tz_text = extract_text(tz_doc.local_path) if tz_doc and tz_doc.local_path else None

    # Считаем simhash от извлечённого текста и кэшируем в documents — это
    # одноразовая операция (≤30мс), пригодится и при повторных проходах.
    tz_simhash_unsigned: int | None = None
    if tz_doc is not None and tz_text is not None:
        if tz_doc.text_simhash is None:
            tz_simhash_unsigned = compute_simhash(tz_text)
            tz_doc.text_simhash = to_signed64(tz_simhash_unsigned)
            session.flush()
        else:
            tz_simhash_unsigned = from_signed64(tz_doc.text_simhash)

    # Дедупликация: если уже есть проанализированный лот с тем же
    # `analyzer_version` и Хэмминговым расстоянием ≤ 3 — копируем результат
    # без LLM. force=True обходит и эту дедупликацию (ручной reanalyze в UI).
    if not force and tz_simhash_unsigned is not None:
        twin = _find_simhash_twin(session, tz_simhash_unsigned, exclude_lot_id=lot.id)
        if twin is not None:
            source, distance = twin
            _copy_analysis(
                session,
                lot,
                source=source,
                tz_sha=tz_sha,
                tz_doc_id=tz_doc.id if tz_doc else None,
                prev=prev,
            )
            log.info(
                "lot %s: reused analysis from lot %s (hamming=%d)",
                lot.id, source.lot_id, distance,
            )
            return True

    # Rule-based классификатор — если правила уверены (confidence ≥ THRESHOLD),
    # пишем результат без LLM. tz_summary остаётся NULL (его заполнит этап 4
    # extractive-summary, либо ручной LLM-reanalyze из UI).
    if not force and not prev_is_llm:
        rule_res = classify_lot(
            lot_name=lot.name,
            lot_extra=lot.extra,
            it_category=_it_subcategory(lot.enstru, lot.name),
            plan_amount=lot.plan_amount,
            announcement_attributes=announcement.attributes if announcement else None,
            tz_text=tz_text,
        )
        if rule_res.confidence >= CONFIDENCE_THRESHOLD:
            if prev is None:
                prev = LotAnalysis(lot_id=lot.id, analyzer_version=RULES_VERSION)
                session.add(prev)
            prev.dev_category = rule_res.dev_category
            prev.tech_stack = list(rule_res.tech_stack)
            prev.tz_summary = extract_summary(
                tz_text=tz_text, lot_name=lot.name, lot_extra=lot.extra
            )
            prev.solo_feasible = rule_res.solo_feasible
            prev.vendor_lock_risk = rule_res.vendor_lock_risk
            prev.analysis_confidence = rule_res.analysis_confidence
            prev.analyzed_at = datetime.now(UTC)
            prev.analyzer_version = RULES_VERSION
            prev.tz_sha256 = tz_sha
            prev.source_document_id = tz_doc.id if tz_doc else None
            prev.reused_from_lot_id = None
            prev.error = None
            session.flush()
            log.info(
                "lot %s: rule-based classify (category=%s, conf=%.2f)",
                lot.id, rule_res.dev_category, rule_res.confidence,
            )
            return True

    outcome = _call_llm(lot, announcement, tz_text)
    # Учитываем токены, даже если результат не распарсился: вызов был платный.
    record_call(session, "analyze", outcome.usage, lot_id=lot.id)
    if outcome.result is None:
        # Не создаём запись на ошибку — иначе при следующем прогоне с
        # работающим ключом мы её не переанализируем (идемпотентность).
        log.warning("LLM skip for lot %s: %s", lot.id, outcome.error)
        return False

    r = outcome.result
    if prev is None:
        prev = LotAnalysis(lot_id=lot.id, analyzer_version=ANALYZER_VERSION)
        session.add(prev)
    prev.dev_category = r.dev_category
    prev.tech_stack = list(r.tech_stack)
    prev.tz_summary = r.tz_summary
    prev.solo_feasible = r.solo_feasible
    prev.vendor_lock_risk = r.vendor_lock_risk
    prev.analysis_confidence = r.analysis_confidence
    prev.analyzed_at = datetime.now(UTC)
    prev.analyzer_version = ANALYZER_VERSION
    prev.tz_sha256 = tz_sha
    prev.source_document_id = tz_doc.id if tz_doc else None
    prev.reused_from_lot_id = None  # этот анализ — свежий LLM-вызов
    prev.error = None
    session.flush()
    return True


def _find_simhash_twin(
    session: Session, target_unsigned: int, *, exclude_lot_id: int
) -> tuple[LotAnalysis, int] | None:
    """Ищет ближайший LotAnalysis с Хэмминговым расстоянием ≤ THRESHOLD.

    Линейный скан по индексу анализов: 70К записей × XOR-popcount ≈ десятки мс.
    LSH-banding не нужен при текущем объёме. Если когда-нибудь упрёмся —
    добавим MinHash banding или GIN по битам.
    """
    candidates = session.execute(
        select(LotAnalysis, Document.text_simhash)
        .join(Document, LotAnalysis.source_document_id == Document.id)
        .where(
            LotAnalysis.analyzer_version == ANALYZER_VERSION,
            LotAnalysis.lot_id != exclude_lot_id,
            LotAnalysis.reused_from_lot_id.is_(None),  # копируем только из «первоисточников»
            Document.text_simhash.is_not(None),
        )
    ).all()
    best: LotAnalysis | None = None
    best_dist = HAMMING_THRESHOLD + 1
    for analysis, sh_signed in candidates:
        d = hamming(target_unsigned, from_signed64(sh_signed))
        if d < best_dist:
            best, best_dist = analysis, d
            if d == 0:
                break  # точный клон — лучше не найдём
    if best is None:
        return None
    return best, best_dist


def _copy_analysis(
    session: Session,
    lot: Lot,
    *,
    source: LotAnalysis,
    tz_sha: str | None,
    tz_doc_id: int | None,
    prev: LotAnalysis | None,
) -> None:
    if prev is None:
        prev = LotAnalysis(lot_id=lot.id, analyzer_version=ANALYZER_VERSION)
        session.add(prev)
    prev.dev_category = source.dev_category
    prev.tech_stack = list(source.tech_stack or [])
    prev.tz_summary = source.tz_summary
    prev.solo_feasible = source.solo_feasible
    prev.vendor_lock_risk = source.vendor_lock_risk
    prev.analysis_confidence = source.analysis_confidence
    prev.analyzed_at = datetime.now(UTC)
    prev.analyzer_version = ANALYZER_VERSION
    prev.tz_sha256 = tz_sha
    prev.source_document_id = tz_doc_id
    prev.reused_from_lot_id = source.lot_id
    prev.error = None
    session.flush()


# === Чат по ТЗ ===

# Кешируем извлечённый текст ТЗ, чтобы при чате на странице лота не
# переоткрывать PDF на каждое сообщение. Файлы на диске иммутабельны
# (downloaded once), поэтому ключ — путь.
@lru_cache(maxsize=128)
def _cached_tz_text(local_path: str) -> str | None:
    return extract_text(local_path)


CHAT_SYSTEM_PROMPT = """Ты — аналитик тендеров, работаешь на стороне потенциального участника
гос.закупок Казахстана. Тебе даны данные лота, поля объявления и полный текст
технической спецификации (ТЗ). Пользователь решает, участвовать ли в закупке
и как подготовить заявку — помогай ему добраться до сути быстро.

Как отвечать:
- Сначала прямой ответ на вопрос в 1-2 предложениях, потом детали и
  обоснование. Не начинай с пересказа вопроса или вводных фраз.
- Опирайся на ТЗ и данные лота. Если ответа в них нет — скажи это прямо
  одной фразой и подскажи, где искать (объявление на goszakup, конкурсная
  документация, проект договора, вопрос организатору), не выдумывай.
- Короткие цитаты из ТЗ давай дословно в кавычках с указанием, откуда
  (раздел, пункт, таблица), когда это подтверждает ответ. Длинные куски не
  копируй — пересказывай.
- Цифры (объёмы, суммы, сроки, проценты) приводи точно, как в документе.
- Если в ТЗ есть противоречие или признак заточки под конкретного поставщика
  и это касается вопроса — отметь.
- Форматирование: обычный текст, БЕЗ markdown — никаких **звёздочек**,
  заголовков с #, таблиц и обратных кавычек. Перечисления — с новой строки
  через тире «— ». Абзацы — пустой строкой.
- Отвечай по-русски, по делу, без воды. Не повторяй уже сказанное в
  предыдущих ответах.

Если пользователь просит разобрать лот целиком («суть задачи», «что здесь
нужно»), дай структурированный разбор в таком порядке: что закупают и в
каком объёме; что именно предстоит сделать исполнителю; требования к
участнику (опыт, лицензии, персонал, обеспечение); сроки и условия оплаты;
риски и подводные камни; какие вопросы стоит задать заказчику. Объём —
8-15 строк, каждый блок — одна-три строки.
"""

_ALMATY = timezone(timedelta(hours=5))


def _fmt_local(dt: datetime | None) -> str | None:
    """UTC из БД → алматинское время строкой; пользователь мыслит местным."""
    if dt is None:
        return None
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=UTC)
    return dt.astimezone(_ALMATY).strftime("%d.%m.%Y %H:%M") + " (Алматы)"


def _plan_point_of(lot: Lot) -> PlanPoint | None:
    # Пункт плана — единственный источник аванса, срока поставки и источника
    # финансирования (правило #26). chat_about_lot сессию не получает —
    # берём ту, к которой привязан lot; отвязанный лот (тесты) → без плана.
    db = object_session(lot)
    if db is None:
        return None
    root = lot.plan_root_id or plan_root_from_number(lot.number)
    return db.get(PlanPoint, root) if root else None


def _chat_system_message(lot: Lot) -> str:
    announcement = lot.announcement
    documents: list[Document] = list(announcement.documents) if announcement else []
    tz_doc = pick_tz_document(documents)
    tz_text = (
        _cached_tz_text(tz_doc.local_path)
        if tz_doc and tz_doc.local_path
        else None
    )

    lines = [
        CHAT_SYSTEM_PROMPT.rstrip(),
        "",
        "--- ДАННЫЕ ЛОТА ---",
        f"Название: {lot.name or '—'}",
        f"ENSTRU: {lot.enstru or '—'}",
        f"Заказчик: {lot.customer.name if lot.customer else '—'}",
        f"Регион: {region_name(lot.kato) or lot.kato or '—'}",
        f"Способ закупки: {lot.method or (announcement.method if announcement else '—')}",
        f"Плановая сумма ₸: {lot.plan_amount or '—'}",
        f"Статус: {lot.status_name or '—'}",
        f"IT-категория (pre-filter): {_it_subcategory(lot.enstru, lot.name) or '—'}",
    ]
    if lot.quantity:
        lines.append(f"Количество: {lot.quantity} {lot.unit or ''}".rstrip())
    if announcement:
        if announcement.application_start:
            lines.append(f"Начало приёма заявок: {_fmt_local(announcement.application_start)}")
        if announcement.application_end:
            lines.append(f"Окончание приёма заявок: {_fmt_local(announcement.application_end)}")
        if announcement.attributes:
            lines.append(f"Признаки: {announcement.attributes}")
    if lot.extra:
        lines.append(f"Доп. характеристика: {lot.extra}")
    if lot.bids_count is not None:
        # Заявки видны только после дедлайна (правило #22) — до него это
        # число ничего не говорит, и модель не должна на него опираться.
        lines.append(f"Подано заявок (видно после дедлайна): {lot.bids_count}")
    if lot.winner_name:
        lines.append(f"Победитель: {lot.winner_name}")

    plan = _plan_point_of(lot)
    if plan:
        plan_lines = []
        if plan.prepayment is not None:
            plan_lines.append(f"аванс {plan.prepayment:g}%")
        if plan.supply_date:
            plan_lines.append(f"срок поставки по плану: {plan.supply_date}")
        if plan.finsource:
            plan_lines.append(f"источник финансирования: {plan.finsource}")
        if plan.month:
            plan_lines.append(f"плановый месяц закупки: {plan.month}")
        if plan_lines:
            lines.append("Из годового плана закупок: " + "; ".join(plan_lines))

    analysis = lot.analysis
    if analysis and analysis.tz_summary:
        lines += ["", "--- КРАТКИЙ БРИФ (авто-анализ) ---", analysis.tz_summary]

    if tz_text:
        lines += [
            "",
            f"--- ТЕХНИЧЕСКАЯ СПЕЦИФИКАЦИЯ ({tz_doc.name}) ---",
            tz_text,
        ]
    else:
        lines += ["", "ТЗ не доступно (нет скачанного файла или нет текстового слоя)."]
    return "\n".join(lines)


def chat_about_lot(lot: Lot, history: list[dict]) -> tuple[str, dict | None]:
    """Один поворот чата: возвращает (ответ ассистента, usage-токены).

    history — список сообщений вида [{"role": "user"|"assistant", "content": str}, ...].
    Последнее сообщение должно быть от user (это и есть текущий вопрос).
    Бросает RuntimeError при проблемах с API/ключом. usage может быть None
    (старый SDK без поля) — учёт расхода тогда просто пропускает этот вызов.
    """
    if not history:
        raise RuntimeError("пустая история")

    api_key = os.environ.get("CEREBRAS_API_KEY")
    if not api_key:
        raise RuntimeError("CEREBRAS_API_KEY не задан")

    try:
        from cerebras.cloud.sdk import Cerebras
    except ImportError as e:
        raise RuntimeError(f"cerebras-cloud-sdk не установлен: {e}") from e

    model = os.environ.get("GZ_LLM_MODEL", DEFAULT_MODEL)
    client = Cerebras(api_key=api_key)

    messages = [{"role": "system", "content": _chat_system_message(lot)}]
    messages.extend(history)

    resp = client.chat.completions.create(
        model=model,
        messages=messages,
        max_tokens=2048,
        # Для чата — medium: вопрос по 10-30 страницам ТЗ требует сопоставить
        # разделы, на low модель отвечала поверхностно. Вызовов единицы в
        # месяц, цена не аргумент.
        reasoning_effort="medium",
        timeout=120.0,
    )
    if not resp.choices:
        raise RuntimeError("пустой ответ модели")
    content = resp.choices[0].message.content
    if not content:
        raise RuntimeError("модель вернула пустой content")
    return content, usage_from_response(resp, model)


# Однокнопочное «Подробнее» из Telegram-уведомления: тот же чат-контекст
# (данные лота + полный текст ТЗ), но фиксированный вопрос. Требования к длине
# заданы в промпте — уведомление читают с телефона, простыня не нужна.
EXPLAIN_USER_PROMPT = (
    "Объясни простым, понятным неспециалисту языком, что это за лот: что "
    "заказчик хочет получить, какие работы предстоят исполнителю и есть ли "
    "важные требования, сроки или подводные камни. Пиши обычным текстом без "
    "markdown, списков и заголовков. Объём — 5-8 предложений: не слишком "
    "коротко и не слишком длинно. Не цитируй ТЗ дословно."
)


def explain_lot(lot: Lot) -> tuple[str, dict | None]:
    """Объяснение лота простым языком (для Telegram). Возвращает (текст, usage).

    Бросает RuntimeError при проблемах с API/ключом — вызывающий actor сам
    решает, что отправить пользователю в этом случае.
    """
    return chat_about_lot(lot, [{"role": "user", "content": EXPLAIN_USER_PROMPT}])


# Машинные коды → русские подписи для UI.
DEV_CATEGORY_LABELS: dict[str, str] = {
    "1c_development": "1С: разработка",
    "web_development": "Веб-разработка",
    "mobile_dev": "Мобильное приложение",
    "integration": "Интеграция систем",
    "software_support": "Поддержка ПО / лицензии",
    "it_infra": "ИТ-инфраструктура",
    "hardware": "Поставка железа",
    "not_dev": "Не разработка",
}

VENDOR_LOCK_LABELS: dict[str, str] = {
    "low": "низкий",
    "medium": "средний",
    "high": "высокий",
}


def dev_category_label(code: str | None) -> str:
    return DEV_CATEGORY_LABELS.get(code or "", code or "—")


def vendor_lock_label(code: str | None) -> str:
    return VENDOR_LOCK_LABELS.get(code or "", code or "—")
